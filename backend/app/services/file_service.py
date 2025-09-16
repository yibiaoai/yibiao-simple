"""文件处理服务"""
import aiofiles
import os
import time
import gc
from datetime import datetime
from typing import Optional, List
import PyPDF2
import docx
from fastapi import UploadFile
from ..config import settings

# 新增的第三方库
try:
    import pdfplumber
    import fitz  # PyMuPDF
    from docx2python import docx2python
    HAS_ADVANCED_LIBS = True
except ImportError as e:
    HAS_ADVANCED_LIBS = False
    print(f"高级文档处理库未安装: {e}")


class FileService:
    """文件处理服务"""
    
    @staticmethod
    def _safe_file_cleanup(file_path: str, max_retries: int = 3) -> bool:
        """安全删除文件，带重试机制"""
        for attempt in range(max_retries):
            try:
                if os.path.exists(file_path):
                    # 强制垃圾回收，释放可能的文件句柄
                    gc.collect()
                    time.sleep(0.1 * (attempt + 1))  # 递增延迟
                    os.remove(file_path)
                return True
            except OSError as e:
                if attempt == max_retries - 1:
                    print(f"无法删除文件 {file_path}: {e}")
                    return False
                time.sleep(0.5)  # 等待后重试
        return True
    
    @staticmethod
    async def save_uploaded_file(file: UploadFile) -> str:
        """保存上传的文件并返回文件路径"""
        # 创建上传目录
        os.makedirs(settings.upload_dir, exist_ok=True)

        # 生成带时间戳的文件名，防止重复
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")[:-3]  # 精确到毫秒
        filename = file.filename or "unknown_file"

        # 分离文件名和扩展名
        name, ext = os.path.splitext(filename)

        # 生成新的文件名：原文件名_时间戳.扩展名
        new_filename = f"{name}_{timestamp}{ext}"
        file_path = os.path.join(settings.upload_dir, new_filename)

        # 异步保存文件
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            await f.write(content)

        return file_path
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """从PDF文件提取文本，支持表格内容"""
        if HAS_ADVANCED_LIBS:
            return FileService._extract_pdf_with_pdfplumber(file_path)
        else:
            # 降级到原来的PyPDF2方法
            return FileService._extract_pdf_with_pypdf2(file_path)
    
    @staticmethod
    def _extract_pdf_with_pdfplumber(file_path: str) -> str:
        """使用pdfplumber提取PDF文本，包含表格（确保及时释放文件句柄）"""
        try:
            extracted_text = []
            
            # 使用上下文管理器，避免在Windows上产生文件锁
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # 添加页码标识
                    extracted_text.append(f"\n--- 第 {page_num} 页 ---\n")
                    
                    # 提取普通文本
                    text = page.extract_text()
                    if text:
                        extracted_text.append(text)
                    
                    # 提取表格
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        extracted_text.append(f"\n[表格 {table_num}]")
                        for row in table:
                            if row:  # 跳过空行
                                # 过滤空值并连接单元格
                                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                extracted_text.append(row_text)
                        extracted_text.append("[表格结束]\n")
            
            result = "\n".join(extracted_text).strip()
            gc.collect()
            return result
        except Exception as e:
            gc.collect()
            # 如果pdfplumber失败，尝试PyMuPDF
            try:
                return FileService._extract_pdf_with_pymupdf(file_path)
            except Exception:
                raise Exception(f"PDF文件读取失败: {str(e)}")
    
    @staticmethod
    def _extract_pdf_with_pymupdf(file_path: str) -> str:
        """使用PyMuPDF提取PDF文本"""
        try:
            doc = fitz.open(file_path)
            extracted_text = []
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                extracted_text.append(f"\n--- 第 {page_num + 1} 页 ---\n")
                
                # 提取文本
                text = page.get_text()
                if text:
                    extracted_text.append(text)
                
                # 尝试提取表格
                try:
                    tables = page.find_tables()
                    for table_num, table in enumerate(tables, 1):
                        extracted_text.append(f"\n[表格 {table_num}]")
                        table_data = table.extract()
                        for row in table_data:
                            if row:
                                row_text = " | ".join([str(cell) if cell else "" for cell in row])
                                extracted_text.append(row_text)
                        extracted_text.append("[表格结束]\n")
                except:
                    # 如果表格提取失败，跳过
                    pass
            
            doc.close()
            return "\n".join(extracted_text).strip()
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")
    
    @staticmethod 
    def _extract_pdf_with_pypdf2(file_path: str) -> str:
        """使用PyPDF2提取PDF文本（原方法）"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text.strip()
        except Exception as e:
            raise Exception(f"PDF文件读取失败: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """从Word文档提取文本，支持表格内容"""
        if HAS_ADVANCED_LIBS:
            return FileService._extract_docx_with_docx2python(file_path)
        else:
            # 降级到原来的python-docx方法，但增强表格处理
            return FileService._extract_docx_with_python_docx(file_path)
    
    @staticmethod
    def _extract_docx_with_docx2python(file_path: str) -> str:
        """使用docx2python提取Word文档内容（确保及时释放文件句柄）"""
        try:
            extracted_text = []
            
            # 使用上下文管理器确保文件及时关闭，避免Windows上的锁定
            with docx2python(file_path) as content:
                # 处理文档内容
                if hasattr(content, 'document'):
                    for section in content.document:
                        for element in section:
                            if isinstance(element, list):
                                # 这可能是表格
                                extracted_text.append("\n[表格内容]")
                                for row in element:
                                    if isinstance(row, list):
                                        row_text = " | ".join([str(cell).strip() for cell in row if cell])
                                        if row_text:
                                            extracted_text.append(row_text)
                                    else:
                                        extracted_text.append(str(row))
                                extracted_text.append("[表格结束]\n")
                            else:
                                # 普通文本
                                text = str(element).strip()
                                if text:
                                    extracted_text.append(text)
            
            result = "\n".join(extracted_text).strip()
            gc.collect()
            return result
        except Exception as e:
            gc.collect()
            # 如果docx2python失败，回退到增强的python-docx
            try:
                return FileService._extract_docx_with_python_docx(file_path)
            except Exception:
                raise Exception(f"Word文档读取失败: {str(e)}")
    
    @staticmethod
    def _extract_docx_with_python_docx(file_path: str) -> str:
        """使用python-docx提取Word文档内容（增强版）"""
        doc = None
        try:
            doc = docx.Document(file_path)
            extracted_text = []
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    extracted_text.append(text)
            
            # 提取表格内容
            for table_num, table in enumerate(doc.tables, 1):
                extracted_text.append(f"\n[表格 {table_num}]")
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text if cell_text else "")
                    row_text = " | ".join(row_data)
                    if row_text.strip():
                        extracted_text.append(row_text)
                extracted_text.append("[表格结束]\n")
            
            result = "\n".join(extracted_text).strip()
            
            # 确保释放资源
            if doc:
                del doc
            gc.collect()
            
            return result
        except Exception as e:
            # 确保释放资源
            if doc:
                del doc
            gc.collect()
            raise Exception(f"Word文档读取失败: {str(e)}")
    
    @staticmethod
    async def process_uploaded_file(file: UploadFile) -> str:
        """处理上传的文件并提取文本内容"""
        # 检查文件大小
        content = await file.read()
        if len(content) > settings.max_file_size:
            raise Exception(f"文件大小超过限制 ({settings.max_file_size / 1024 / 1024}MB)")
        
        # 重置文件指针
        await file.seek(0)
        
        # 保存文件
        file_path = await FileService.save_uploaded_file(file)
        
        try:
            # 根据文件类型提取文本
            if file.content_type == "application/pdf":
                text = FileService.extract_text_from_pdf(file_path)
            elif file.content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                text = FileService.extract_text_from_docx(file_path)
            else:
                raise Exception("不支持的文件类型，请上传PDF或Word文档")
            
            # 成功提取后，使用安全的文件清理方法
            FileService._safe_file_cleanup(file_path)
            
            return text
            
        except Exception as e:
            # 异常情况下也使用安全的文件清理方法
            FileService._safe_file_cleanup(file_path)
            raise e