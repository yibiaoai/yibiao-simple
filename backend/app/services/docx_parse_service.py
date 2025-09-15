"""DOCX 结构与样式解析服务

使用 python-docx 提取：
- 章节层级（基于样式名与编号层级的启发式）
- 段落文本、样式、对齐、加粗/斜体/下划线存在性
- 表格（网格文本、行列数、样式名）
- 图片（数量、尺寸、类型）
- 节（页眉/页脚文本、页面设置、边距）
- 文档样式清单（名称、类型、基类）

注：本模块仅做“读取/建模”，不修改源文档。
"""

from __future__ import annotations

from typing import Any, Dict, List, Optional

import docx
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocxParseService:
    """提供对 DOCX 的结构化解析能力。"""

    @staticmethod
    def parse(file_path: str) -> Dict[str, Any]:
        """解析 DOCX 文件为结构化数据。

        返回的数据包含：styles、paragraphs、headings、tables、images、sections、style_map。
        """
        document = docx.Document(file_path)

        styles = DocxParseService._collect_styles(document)
        paragraphs, headings = DocxParseService._collect_paragraphs(document)
        tables = DocxParseService._collect_tables(document)
        images = DocxParseService._collect_images(document)
        sections = DocxParseService._collect_sections(document)

        # 生成简单样式映射（默认同名映射，后续可由前端/用户调整）
        used_style_names = sorted({p["style_name"] for p in paragraphs if p.get("style_name")})
        style_map = {name: name for name in used_style_names}

        return {
            "styles": styles,
            "paragraphs": paragraphs,
            "headings": headings,
            "tables": tables,
            "images": images,
            "sections": sections,
            "style_map": style_map,
        }

    @staticmethod
    def _collect_styles(document: docx.Document) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        for style in document.styles:
            try:
                style_type = DocxParseService._style_type_name(style.type)
            except Exception:
                style_type = "unknown"
            base = None
            try:
                base = style.base_style.name if style.base_style is not None else None
            except Exception:
                base = None
            results.append({
                "name": getattr(style, "name", None),
                "type": style_type,
                "based_on": base,
            })
        return results

    @staticmethod
    def _style_type_name(style_type: Optional[WD_STYLE_TYPE]) -> str:
        if style_type is None:
            return "unknown"
        mapping = {
            WD_STYLE_TYPE.PARAGRAPH: "paragraph",
            WD_STYLE_TYPE.CHARACTER: "character",
            WD_STYLE_TYPE.TABLE: "table",
            WD_STYLE_TYPE.LIST: "list",
        }
        return mapping.get(style_type, "unknown")

    @staticmethod
    def _collect_paragraphs(document: docx.Document) -> tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        paragraphs: List[Dict[str, Any]] = []
        headings: List[Dict[str, Any]] = []
        for idx, p in enumerate(document.paragraphs):
            text = (p.text or "").strip()
            style_name = getattr(p.style, "name", None)
            alignment = DocxParseService._alignment_name(p.alignment)
            has_bold = any(run.bold for run in p.runs if run.text)
            has_italic = any(run.italic for run in p.runs if run.text)
            has_underline = any(run.underline for run in p.runs if run.text)

            num_id, ilvl = DocxParseService._get_numbering(p)
            is_heading, level = DocxParseService._get_heading_level(style_name)

            para_info = {
                "index": idx,
                "text": text,
                "style_name": style_name,
                "alignment": alignment,
                "has_bold": has_bold,
                "has_italic": has_italic,
                "has_underline": has_underline,
                "num_id": num_id,
                "ilvl": ilvl,
                "is_heading": is_heading,
                "heading_level": level,
            }
            paragraphs.append(para_info)

            if is_heading:
                headings.append({
                    "index": idx,
                    "text": text,
                    "level": level,
                    "style_name": style_name,
                })

        return paragraphs, headings

    @staticmethod
    def _alignment_name(alignment: Optional[WD_ALIGN_PARAGRAPH]) -> Optional[str]:
        if alignment is None:
            return None
        mapping = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
            WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
        }
        return mapping.get(alignment, None)

    @staticmethod
    def _get_numbering(paragraph: docx.text.paragraph.Paragraph) -> tuple[Optional[int], Optional[int]]:
        """读取段落的编号信息（numId 与 ilvl）。若无则返回 (None, None)。"""
        try:
            pPr = paragraph._p.pPr  # type: ignore[attr-defined]
            if pPr is None or pPr.numPr is None:
                return None, None
            num_id_elm = pPr.numPr.numId
            ilvl_elm = pPr.numPr.ilvl
            num_id = int(num_id_elm.val) if num_id_elm is not None else None
            ilvl = int(ilvl_elm.val) if ilvl_elm is not None else None
            return num_id, ilvl
        except Exception:
            return None, None

    @staticmethod
    def _get_heading_level(style_name: Optional[str]) -> tuple[bool, Optional[int]]:
        """根据样式名推断标题级别。兼容中英文字样式，如“Heading 1”、“标题 1”。"""
        if not style_name:
            return False, None
        name = style_name.strip().lower()
        # 英文 Heading
        if name.startswith("heading "):
            try:
                level = int(name.replace("heading ", "").strip())
                return True, level
            except ValueError:
                return True, None
        # 中文 标题 N
        if name.startswith("标题"):
            tail = name.replace("标题", "").strip()
            try:
                level = int(tail)
                return True, level
            except ValueError:
                return True, None
        return False, None

    @staticmethod
    def _collect_tables(document: docx.Document) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        for t_idx, table in enumerate(document.tables):
            style_name = None
            try:
                style_name = table.style.name if table.style is not None else None
            except Exception:
                style_name = None
            grid: List[List[str]] = []
            for row in table.rows:
                row_data: List[str] = []
                for cell in row.cells:
                    row_data.append((cell.text or "").strip())
                grid.append(row_data)
            results.append({
                "index": t_idx,
                "rows": len(table.rows),
                "cols": len(table.columns) if hasattr(table, "columns") else (len(grid[0]) if grid else 0),
                "style_name": style_name,
                "grid": grid,
            })
        return results

    @staticmethod
    def _collect_images(document: docx.Document) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        try:
            for i_idx, shape in enumerate(document.inline_shapes):
                width = getattr(shape, "width", None)
                height = getattr(shape, "height", None)
                shape_type = getattr(shape, "type", None)
                results.append({
                    "index": i_idx,
                    "width": int(width) if width is not None else None,
                    "height": int(height) if height is not None else None,
                    "shape_type": str(shape_type) if shape_type is not None else None,
                })
        except Exception:
            # 某些文档不存在 inline_shapes 或解析失败，忽略即可
            pass
        return results

    @staticmethod
    def _collect_sections(document: docx.Document) -> List[Dict[str, Any]]:
        results: List[Dict[str, Any]] = []
        for s_idx, section in enumerate(document.sections):
            # 页眉/页脚文本汇总
            header_text = DocxParseService._header_footer_text(section.header)
            footer_text = DocxParseService._header_footer_text(section.footer)
            page_width = int(getattr(section.page_width, "_EMUS", section.page_width)) if hasattr(section.page_width, "_EMUS") else int(section.page_width)
            page_height = int(getattr(section.page_height, "_EMUS", section.page_height)) if hasattr(section.page_height, "_EMUS") else int(section.page_height)
            margins = {
                "top": int(section.top_margin),
                "bottom": int(section.bottom_margin),
                "left": int(section.left_margin),
                "right": int(section.right_margin),
                "header": int(section.header_distance),
                "footer": int(section.footer_distance),
            }
            results.append({
                "index": s_idx,
                "page_size": {"width": page_width, "height": page_height},
                "margins": margins,
                "header_text": header_text,
                "footer_text": footer_text,
            })
        return results

    @staticmethod
    def _header_footer_text(hf_part: Any) -> str:
        try:
            lines: List[str] = []
            for p in hf_part.paragraphs:
                t = (p.text or "").strip()
                if t:
                    lines.append(t)
            return "\n".join(lines)
        except Exception:
            return ""


